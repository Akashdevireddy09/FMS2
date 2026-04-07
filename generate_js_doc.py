#!/usr/bin/env python3
"""
Generate a Word document containing all JavaScript files from the FMS project.
Organized with clear file structure and readable formatting.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

# Define all JS files to include
JS_FILES = [
    "app.js",
    "js/app.bootstrap.js",
    "js/core/storage.js",
    "js/core/security.js",
    "js/core/ui.js",
    "js/core/pricing.js",
    "js/modules/authentication.module.js",
    "js/modules/user-dashboard.module.js",
    "js/modules/admin-dashboard-flight-management.module.js",
    "js/modules/available-flights.module.js",
    "js/modules/flight-search.module.js",
    "js/modules/flight-selection.module.js",
    "js/modules/preview.module.js",
    "js/modules/fare-discount.module.js",
    "js/modules/payment.module.js",
    "js/modules/booking-confirmation.module.js",
    "js/modules/my-bookings.module.js",
    "js/modules/cancellation-refund.module.js",
]

def create_js_documentation():
    """Create a Word document with all JS files."""
    doc = Document()
    
    # Add title
    title = doc.add_heading("FMS Project - JavaScript Code Documentation", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run("March 29, 2026\n")
    meta.add_run("Project: ").bold = True
    meta.add_run("Flight Management System (FMS)\n")
    meta.add_run("Total Files: ").bold = True
    meta.add_run(f"{len(JS_FILES)}")
    
    # Add table of contents
    doc.add_heading("Table of Contents", level=2)
    toc_list = doc.add_paragraph()
    for idx, file_path in enumerate(JS_FILES, 1):
        p = doc.add_paragraph(f"{idx}. {file_path}", style="List Number")
    
    doc.add_page_break()
    
    # Add each JS file
    for idx, file_path in enumerate(JS_FILES, 1):
        full_path = Path(file_path)
        
        if not full_path.exists():
            print(f"Warning: File not found - {file_path}")
            continue
        
        # File heading
        heading = doc.add_heading(f"{idx}. {file_path}", level=2)
        
        # File info
        file_info = doc.add_paragraph()
        file_info.add_run("File Path: ").bold = True
        file_info.add_run(f"{file_path}\n")
        file_info.add_run("Type: ").bold = True
        file_info.add_run("JavaScript Module\n")
        
        try:
            with open(full_path, 'r', encoding='utf-8') as f:
                code_content = f.read()
            
            # Add code with monospace font
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(code_content)
            
            # Format code with monospace font
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            code_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add a separator
            doc.add_paragraph("_" * 80)
            doc.add_page_break()
            
        except Exception as e:
            error_para = doc.add_paragraph(f"Error reading file: {str(e)}")
            error_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # Save document
    output_path = "FMS_JavaScript_Code.docx"
    doc.save(output_path)
    
    print(f"✓ Word document created successfully: {output_path}")
    print(f"✓ Total files included: {len(JS_FILES)}")
    print(f"✓ Document size: {os.path.getsize(output_path) / 1024:.2f} KB")
    return output_path

if __name__ == "__main__":
    create_js_documentation()
